import streamlit as st
import pandas as pd
from fuzzywuzzy import process
from datetime import datetime
from io import StringIO
from docx import Document

# 定义题目和答案
questions_answers = {
    "1. 在光滑斜面上，放有物体X和物体Y，一根轻质绳将两个物体连接在一起。如图所示，将沿着斜面向上的力F作用在物体Y上。如果斜面倾角变大，绳的张力会有怎样的变化？A.不变 B.变小 C.变大": 
    """在θ不变的情况下对两个物块整体沿着斜面方向进行受力分析，再根据牛顿第二定律算出两个物块沿着斜面的共同加速度 F - (mX + mY)gsinθ = (mX + mY)a；
然后再用隔离法研究物块X的受力情况，根据牛顿第二定律列出方程 F - mXgsinθ = mX a；带入共同加速度 a 可以算出绳的张力 F = mX / (mX + mY) * a。""",
    
    "2. 一木块置于光滑的水平桌面上，一根轻质绳绕过滑轮连接在木块上。一名男子用大小为10N的力向下拉动绳子，使得木块由静止开始向右运动（如左图），此时木块运动的加速度为 a左。现用质量为1kg的铁球与绳的下端连接，仍然使木块从静止开始向右运动（如右图），此时木块运动的加速度为 a右。问：a左与a右之间的大小关系是怎样的？A.a左=a右 B.a左>a右 C.a左<a右": 
    """对于左图，木块所受的合外力为10N，使木块产生加速度；
而右图的情况则不同，与木块相连的铁球自身也具有与木块加速度大小相同方向向下的加速度，由牛顿第二定律可知，其所受的重力大于绳子的拉力，则右图木块所受的合外力（等于绳的拉力）小于10N，因此右图中木块的加速度小于左图中木块的加速度。""",
    
    "3. 质量分别为2m、m的两个物体在光滑的水平面上相互接触。水平向右的力F施加在质量为2m的物体上（如左图）。现将力F作用在质量为m的物体上（如右图）。与左图相比，右图中两个物体之间的作用力会发生怎样的变化？A.保持不变 B.变大 C.变小": 
    """因为外力F大小不变，可以看出两系统的加速度大小相同，在左图中外力F只直接作用在质量为2m的物块上，没有直接作用在质量为m的物块上，所以质量为m的物块的加速度由质量为2m的物块对它的作用力提供，由此可以求出左图中两物块的相互作用力为 F / 3；
同理右图中两物体间的相互作用力为 2F / 3，所以右图中两物体之间的作用力大于左图中两物体间的作用力。"""
}

# 创建一个用于存储学生搜索结果的 DataFrame
def load_results(student_id):
    file_path = "results_{}.csv".format(student_id)
    try:
        return pd.read_csv(file_path)
    except FileNotFoundError:
        return pd.DataFrame(columns=["时间", "题目", "答案"])

def save_results(results, student_id):
    file_path = "results_{}.csv".format(student_id)
    results.to_csv(file_path, index=False)

def to_csv(df):
    """将 DataFrame 转换为 CSV 格式的字符串"""
    return df.to_csv(index=False)

# Streamlit 应用
st.title('问题与答案查询系统')

# 输入学生 ID
student_id = st.text_input("请输入你的学生 ID:", "")

if student_id:
    # 输入题目
    question_input = st.text_input("请输入题目内容:")

    if question_input:
        # 使用 fuzzywuzzy 库找到最匹配的题目
        best_match, score = process.extractOne(question_input, questions_answers.keys())
        
        # 设置一个匹配阈值
        if score > 80:  # 可以根据需要调整阈值
            answer = questions_answers[best_match]
        else:
            answer = "没有找到对应的答案"
        
        # 使用传统字符串格式化
        st.markdown("**对应的答案是**:\n{}".format(answer.replace(';', ';\n')))
        
        # 获取当前时间
        current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # 读取现有结果
        results = load_results(student_id)
        
        # 添加新的结果
        new_result = {"时间": current_time, "题目": question_input, "答案": answer}
        results = results.append(new_result, ignore_index=True)
        
        # 保存到文件
        save_results(results, student_id)

        # 显示导出按钮
        if st.button("导出结果到 Word 文档"):
            # 生成 CSV 内容
            csv_data = to_csv(results)
            
            # 使用 StringIO 创建一个文件流
            csv_buffer = StringIO(csv_data)
            
            # 文件保存路径
            export_path = r"D:\台式机桌面内容备份\Desktop\ChatGPT\+生成式人工智能\搜题app\results_{}.docx".format(student_id)
            
            # 将 CSV 内容写入 Word 文档
            doc = Document()
            doc.add_heading(f'学生 {student_id} 的查询结果', 0)
            
            # 添加表格到文档
            if not results.empty:
                table = doc.add_table(rows=results.shape[0] + 1, cols=results.shape[1])
                hdr_cells = table.rows[0].cells
                for i, column in enumerate(results.columns):
                    hdr_cells[i].text = column
                
                for row_idx, row in results.iterrows():
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, value in enumerate(row):
                        row_cells[col_idx].text = str(value)
            
            # 保存到指定路径
            doc.save(export_path)
            
            # 提供下载链接
            with open(export_path, "rb") as f:
                st.download_button(
                    label="下载 Word 文档",
                    data=f,
                    file_name="results_{}.docx".format(student_id),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

